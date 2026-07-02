from __future__ import annotations

import html
import os
import zipfile
from pathlib import Path

from docx import Document
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from app.repositories.repository import repo


def get_vietnamese_font() -> str:
    paths = [
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
    ]
    for p in paths:
        if os.path.exists(p):
            try:
                pdfmetrics.registerFont(TTFont('VietnameseFont', p))
                return 'VietnameseFont'
            except:
                pass
    return 'Helvetica'


def export_project(project_id: str, out_path: str | None = None, fmt: str = "txt", from_chapter: int = 1, to_chapter: int | None = None, overwrite: bool = False) -> dict:
    project = repo.get_project(project_id)
    chapters = [c for c in repo.chapters(project_id) if c.get("status") == "committed"]
    if to_chapter is not None:
        chapters = [c for c in chapters if from_chapter <= c["chapter_no"] <= to_chapter]
    else:
        chapters = [c for c in chapters if c["chapter_no"] >= from_chapter]
    if not chapters:
        raise ValueError("Chưa có chương đã hoàn thành để export")
    fmt = fmt.lower().strip(".")
    if not out_path:
        out_path = str(Path(project.get("output_dir") or ".") / f"{project['name']}.{fmt}")
    path = Path(out_path)
    if path.exists() and not overwrite:
        raise FileExistsError(f"File đã tồn tại: {path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    if fmt == "txt":
        body = "\n\n".join(f"{c.get('title') or 'Chương ' + str(c['chapter_no'])}\n\n{c.get('final_text') or ''}" for c in chapters)
        path.write_text(body, encoding="utf-8")
    elif fmt == "epub":
        render_epub(path, project["name"], chapters)
    elif fmt == "docx":
        render_docx(path, project["name"], chapters)
    elif fmt == "pdf":
        render_pdf(path, project["name"], chapters)
    else:
        raise ValueError("Chỉ hỗ trợ các định dạng: txt, epub, docx, pdf")
    return {"path": str(path), "chapters": len(chapters), "bytes": path.stat().st_size}


def render_epub(path: Path, title: str, chapters: list[dict]) -> None:
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        z.writestr("META-INF/container.xml", """<?xml version='1.0'?><container version='1.0' xmlns='urn:oasis:names:tc:opendocument:xmlns:container'><rootfiles><rootfile full-path='OEBPS/content.opf' media-type='application/oebps-package+xml'/></rootfiles></container>""")
        manifest = []
        spine = []
        for c in chapters:
            name = f"chapter-{c['chapter_no']}.xhtml"
            manifest.append(f"<item id='ch{c['chapter_no']}' href='{name}' media-type='application/xhtml+xml'/>")
            spine.append(f"<itemref idref='ch{c['chapter_no']}'/>")
            z.writestr(f"OEBPS/{name}", f"<html xmlns='http://www.w3.org/1999/xhtml'><head><title>{html.escape(c.get('title') or '')}</title></head><body><h1>{html.escape(c.get('title') or '')}</h1><p>{html.escape(c.get('final_text') or '').replace(chr(10), '<br/>')}</p></body></html>")
        z.writestr("OEBPS/content.opf", f"""<?xml version='1.0' encoding='utf-8'?><package version='3.0' xmlns='http://www.idpf.org/2007/opf' unique-identifier='bookid'><metadata xmlns:dc='http://purl.org/dc/elements/1.1/'><dc:title>{html.escape(title)}</dc:title><dc:language>vi</dc:language><dc:identifier id='bookid'>story-clone-{html.escape(title)}</dc:identifier></metadata><manifest>{''.join(manifest)}</manifest><spine>{''.join(spine)}</spine></package>""")


def render_docx(path: Path, title: str, chapters: list[dict]) -> None:
    doc = Document()
    doc.add_heading(title, level=0)
    for c in chapters:
        c_title = c.get('title') or f"Chương {c['chapter_no']}"
        doc.add_heading(c_title, level=1)
        text = c.get('final_text') or ''
        for paragraph in text.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
            else:
                doc.add_paragraph('')
    doc.save(path)


def render_pdf(path: Path, title: str, chapters: list[dict]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    font_name = get_vietnamese_font()
    
    title_style = ParagraphStyle(
        'PDFTitle',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=24,
        leading=28,
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    h1_style = ParagraphStyle(
        'PDFH1',
        parent=styles['Heading2'],
        fontName=font_name,
        fontSize=18,
        leading=22,
        spaceBefore=15,
        spaceAfter=10,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'PDFBody',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=12,
        leading=16,
        spaceAfter=8,
        alignment=TA_JUSTIFY
    )
    
    story.append(Paragraph(html.escape(title), title_style))
    story.append(Spacer(1, 20))
    
    for c in chapters:
        c_title = c.get('title') or f"Chương {c['chapter_no']}"
        story.append(Paragraph(html.escape(c_title), h1_style))
        text = c.get('final_text') or ''
        for para in text.split('\n'):
            para_text = para.strip()
            if para_text:
                story.append(Paragraph(html.escape(para_text), body_style))
        story.append(Spacer(1, 10))
        
    doc.build(story)

